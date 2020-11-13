# This script reads a specific .htm file in current directory and then converts the string of that text into a word docx file

#import docx library
from docx import Document
from docx.shared import Inches

# main function
def readFileAndConvertToWordDoc():
    
    # define files name of file which will be read and converted
    file = 'data.htm'
    
    # with statement to open file, read file and convert to word docx
    with open(file,  'r') as f:
        
        # reads file data
        data = f.read()
        print(data)

        # converts file data into word docx
        document = Document()
        document.add_heading('Output of test.htm', 0)
        p = document.add_paragraph(data)
        document.save('test.docx')
        


# run main function to read file, convert to word docx and save new file
readFileAndConvertToWordDoc()



