from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Test document', 0)

p = document.add_paragraph('This is a test word document')

document.save('test.docx')

